# TrustForge - 講稿生成腳本
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 頁面設定：A4 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── 工具函式 ──
def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color or RGBColor(0x1d, 0x4e, 0xd8)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    # 底線分隔
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '93C5FD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def timing_badge(text):
    """時間提示（橘色粗體）"""
    p = doc.add_paragraph()
    run = p.add_run(f"⏱ {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xea, 0x58, 0x0c)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def speech(text):
    """逐字稿（深色引用框風格）"""
    p = doc.add_paragraph()
    # 左邊框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '3B82F6')
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'「{text}」')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    run.font.italic = True
    return p

def note(text):
    """備註/提示（灰色小字）"""
    p = doc.add_paragraph()
    run = p.add_run(f"📌 {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def bullet(text, bold_prefix=None):
    """條列項目"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run_b.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        run_r = p.add_run(text)
        run_r.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    return p

def qa_item(q, a):
    p = doc.add_paragraph()
    r1 = p.add_run(f"Q：{q}\n")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
    r2 = p.add_run(f"A：{a}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    return p

def divider():
    doc.add_paragraph("─" * 50)

# ══════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════
h1("TrustForge Hermes 信源熔爐")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("六分鐘簡報 + 四分鐘備詢 ── 逐字講稿")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("2026 雲湧智生黑客松 · HOYA BIT 智慧交易命題 · Team 11 中再參與")
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("隊長：王英豪　成員：嵋婕　林子彤　王榆翔 Nicholas")
r3.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  使用說明
# ══════════════════════════════════════════════════════
h2("使用說明", color=RGBColor(0x0f, 0x76, 0x6e))

bullet("藍色引用框 = 對評審說的話（逐字稿）")
bullet("⏱ 橘色 = 時間控制點")
bullet("📌 灰色 = 備忘/操作提示，不說出口")
bullet("Q&A 區 = 備詢四分鐘，熟記後不用照唸")
bullet("第一階段 15 分鐘 Live Demo 是在跑系統，不用這份稿；這份稿用於第二階段 6 分鐘簡報")

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  PART A — 六分鐘簡報逐字稿
# ══════════════════════════════════════════════════════
h1("PART A — 六分鐘簡報逐字稿")

# ── 第一段：開場痛點 ──
h2("【第一段】開場痛點　0:00 – 1:00")
timing_badge("目標：1 分鐘。說完痛點，讓評審有感。")

speech(
    "各位評審好，我們是 Team 11，中再參與，出品 TrustForge。"
)
speech(
    "加密市場每天產生數百萬條資訊。問題不是資訊太少——"
    "問題是你根本不知道哪條該信。"
)
speech(
    "一項對 10 個商業 AI、近七萬條引用的實測顯示，"
    "LLM 的引用幻覺率高達 11.4% 到 56.8%。"
    "也就是說，超過一成、甚至近六成的引用，可能根本不存在。"
    "在這種環境下用 AI 做加密市場決策，風險很高。"
)
note("停頓一秒，讓數字沉澱")
speech(
    "一般 RAG 系統的解法是把所有來源等重丟給 LLM 摘要——"
    "這沒有解決問題，只是讓問題更隱蔽。"
)

# ── 第二段：我們是什麼 ──
h2("【第二段】TrustForge 是什麼　1:00 – 2:00")
timing_badge("目標：1 分鐘。講清楚三層架構與核心差異。")

speech(
    "TrustForge 是一個多源資訊「信任提煉」Agent。"
    "它的核心差異在中間這一層——Trust Layer。"
)
speech(
    "多源資訊進來之後，不直接丟給 LLM。"
    "我們先對每一條主張，用四個維度計算信任分數："
    "來源信譽、交叉佐證、時效衰減、操縱懲罰。"
    "加權之後，才讓 Bedrock 行文。"
)
speech(
    "Bedrock 在這裡只負責把推理「寫成人話」——"
    "判斷結構、證據整合、信任評分，全部是我們自己的 pipeline 產生的。"
    "這就是反作弊設計，也是命題的核心要求。"
)
note("切換到架構圖投影片")

# ── 第三段：四大亮點 ──
h2("【第三段】四大技術亮點　2:00 – 4:00")
timing_badge("目標：2 分鐘。快速走過四武器，配合 Live Demo 畫面。")

speech("接下來我快速展示四個關鍵設計。")

bullet("① 反事實 A/B 對照", "")
speech(
    "同一題，我們同時跑「關掉信任層的 naive RAG」和「TrustForge」。"
    "你可以看到 naive 版被兩條機器人假新聞帶偏，"
    "TrustForge 的操縱旗標介入，把這兩條壓到 0.12 分，移進反方清單。"
)
note("展示 A/B 對照截圖（賽前已備好）")

bullet("② 來源獨立性圖譜", "")
speech(
    "這 14 篇新聞，去重之後只有 3 個真獨立來源。"
    "我們不會因為有 14 篇就說『多方佐證』——"
    "交叉佐證計的是獨立來源數，不是文章數。"
)

bullet("③ 矛盾帳本", "")
speech(
    "當資金費率偏多、但鏈上同時出現大額流出，"
    "TrustForge 不會硬給一個方向——"
    "它會顯示矛盾帳本，告訴你兩個訊號都存在，並說明各自的信任權重。"
)

bullet("④ 負空間情報", "")
speech(
    "這次監管面沒有任何新事件。"
    "一般系統會跳過這個維度，我們不會——"
    "「未觀察到監管訊號」本身就是情報，我們標記出來。"
)
note("配合 Evidence Panel 畫面")

# ── 第四段：Live Demo 結果 ──
h2("【第四段】Live Demo 結果走讀　4:00 – 5:30")
timing_badge("目標：1.5 分鐘。讓評審親眼看到溯源鏈。")

speech(
    "大家看到這份報告——每一個結論都有對應的 Evidence ID。"
    "比如這條 E3，點開來："
    "來源是 CoinGlass API，抓取時間是今天上午，"
    "引用的是 SOL 24 小時交易所淨流出數據，"
    "對應的主張是『鏈上出現潛在賣壓』。"
    "每一步都可以查。"
)
note("點開 evidence.json 的 E3 給評審看")
speech(
    "執行紀錄顯示：總共跑了 11 分 42 秒，"
    "使用了 78% 的 15 分鐘預算，"
    "Bedrock 呼叫三次：claim 抽取、判斷整合、敘事行文。"
    "全程有時戳，可以審計。"
)
note("展示 execution_log.jsonl")

# ── 第五段：結尾 ──
h2("【第五段】誠實結尾　5:30 – 6:00")
timing_badge("目標：30 秒。誠實定位，收尾有力。")

speech(
    "我們誠實說：TrustForge 不預測幣價，不保證勝率。"
    "我們解決的是更根本的問題："
    "讓你看到資訊從哪來、有多完整、有無來源互相矛盾。"
)
speech(
    "Nansen、LunarCrush、Arkham 都給你分數——"
    "但它們都沒有給你這一層：可溯源、可審計的信任鏈。"
    "這就是 TrustForge。謝謝。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  PART B — 四分鐘備詢
# ══════════════════════════════════════════════════════
h1("PART B — 四分鐘備詢 Q&A")

p = doc.add_paragraph()
r = p.add_run("以下為預期評審追問與標準應答。熟記核心邏輯即可，不需照唸。")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
r.font.italic = True

divider()

qa_item(
    "Trust Score 準嗎？能預測漲跌嗎？",
    "我們的定位是「資訊完整度 + 可溯源」，不是價格預測。"
    "內部驗證純預測 AUC 約 0.49，接近隨機，我們誠實自曝這一點。"
    "TrustForge 解決的是「哪條資訊該信」，不是「幣價往哪走」。"
)

qa_item(
    "競品 Nansen / LunarCrush 也有分數，你們差在哪？",
    "他們有分數，但沒有可溯源的證據鏈。"
    "LunarCrush Galaxy Score 是黑箱，你不知道哪條消息支撐了哪個判斷。"
    "TrustForge 補上的是每個結論 → claim_id → 原始來源 URL + 時間戳，全程可審計。"
)

qa_item(
    "如果所有來源都造假、都是機器人怎麼辦？",
    "ManipulationPenalty 的設計：Bot 協同轉發會被模板相似聚類偵測（Jaccard ≥ 0.8 × 3 個來源觸發），"
    "分數壓低並移入反方列表，不靜默丟棄。"
    "讓評審看到「有大量可疑訊號」這個事實本身，就是有價值的情報。"
)

qa_item(
    "為什麼不給明確買賣建議？",
    "HOYA BIT 合規要求，也是命題精神。"
    "我們輸出「可查證的判斷依據」，不代替你決策。"
    "這與 HOYA BIT「AI 輔助決策、不代替決策」的理念完全對齊。"
)

qa_item(
    "模型合規嗎？有沒有用到 OpenAI 或其他服務？",
    "全程只用 AWS Bedrock（Claude Haiku 4.5）。"
    "所有 LLM 呼叫集中在 bedrock.py，競賽帳號直連 bedrock-runtime，"
    "沒有任何第三方 LLM 閘道。"
)

qa_item(
    "HOYA BIT 企業數據怎麼用的？",
    "HOYA BIT 行情資料在 Trust Layer 裡是最高信任來源之一，信譽分 0.85。"
    "5 幣 × 5 年 Daily OHLCV 作為價格事實錨點，每個價格主張都帶 SHA-256 校驗。"
)

qa_item(
    "為什麼用 AWS Kiro？",
    "Kiro 讓我們用 Steering 管理競賽規範、Hooks 自動化品質控管、Spec 生成連接器規格。"
    "整個開發過程有完整紀錄，.kiro 資料夾已包含 specs / hooks / steering，上傳 GitHub 可驗證。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  PART C — 賽前提案交流重點（8/1 上午 10:00-12:00）
# ══════════════════════════════════════════════════════
h1("PART C — 賽前提案交流重點")

p = doc.add_paragraph()
r = p.add_run("8/1（六）10:00–12:00 與業師及評審提案交流，非正式，重點是讓對方了解你們在做什麼。")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

h2("一句話定位")
speech(
    "我們做的是加密市場資訊的信任提煉——"
    "把多源雜訊先評可信度、做交叉佐證，"
    "讓每一個結論都可以溯源到原始資料。"
)

h2("三層架構（30 秒說完）")
bullet("Layer 1：平行抓取 6 類資料來源（價格/鏈上/新聞/社群/HOYA BIT/監管）")
bullet("Layer 2：Trust Layer — 逐條評信任分（核心差異化）")
bullet("Layer 3：AWS Bedrock 行文，只引用已加權的可信摘要")

h2("跟 HOYA BIT 的連結")
bullet("HOYA BIT 行情是最高信任來源（信譽分 0.85）")
bullet("輸出定位是「輔助決策工具」，不給投資建議，符合 HOYA BIT 合規要求")
bullet("商業路徑：Trust Layer API 可白標給其他交易所")

h2("預計展示")
bullet("8/2 下午抽題後，15 分鐘內跑完 + 上傳 4 份交付件")
bullet("第二階段 6 分鐘簡報 + 4 分鐘問答")

# ── 輸出 ──
out_path = os.path.join(os.path.expanduser("~"), "Desktop", "TrustForge_講稿.docx")
doc.save(out_path)
print(f"✅ 已輸出：{out_path}")

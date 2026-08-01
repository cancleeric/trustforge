from docx import Document
from pathlib import Path
from html import escape

src=Path('outputs/TrustForge_完整商業化提案報告.docx')
dst=Path('outputs/TrustForge_完整商業化提案報告.html')
d=Document(src)
paragraphs={p._p:p for p in d.paragraphs}
tables={t._tbl:t for t in d.tables}
parts=['<!doctype html><html lang="zh-Hant"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>TrustForge Hermes｜完整商業化提案報告</title><style>']
parts.append('body{font-family:Calibri,"Microsoft JhengHei",sans-serif;max-width:960px;margin:0 auto;padding:42px 34px 80px;color:#172033;line-height:1.72;background:#fff}h1{font-size:26px;color:#2e74b5;margin:30px 0 12px;border-bottom:2px solid #dbe6ef;padding-bottom:5px}h2{font-size:19px;color:#2e74b5;margin:22px 0 8px}h3{font-size:16px;color:#1f4d78;margin:16px 0 6px}p{font-size:15px;margin:0 0 10px;text-align:justify}.cover{text-align:center;margin:30px 0 28px}.cover p{text-align:center}.cover .brand{font-weight:700;color:#5b6573}.cover h1{font-size:34px;color:#0b2545;border:0;margin:6px 0}.cover .sub{font-size:19px;color:#5b6573}.callout{background:#eef8fb;border-left:5px solid #2e74b5;padding:12px 16px;margin:14px 0}.warn{background:#fff8e8;border-left-color:#7a5a00}ul,ol{margin:6px 0 12px 28px;padding:0}li{margin:4px 0;font-size:15px}table{width:100%;border-collapse:collapse;margin:12px 0 16px;font-size:14px}th,td{border:1px solid #cbd7e3;padding:8px 10px;vertical-align:top}th{background:#e8eef5;color:#0b2545;text-align:center}footer{color:#68778a;border-top:1px solid #d7e3ef;margin-top:36px;padding-top:8px;font-size:12px}@media print{body{padding:20px}h1{break-after:avoid}table, .callout{break-inside:avoid}}</style></head><body>')
first=True; in_cover=False; list_tag=None
def close_list():
    global list_tag
    if list_tag:
        parts.append(f'</{list_tag}>')
        list_tag=None

for child in d.element.body.iterchildren():
    if child in paragraphs:
        p=paragraphs[child]; text=escape(p.text); st=p.style.name
        if not text: continue
        if first:
            parts.append(f'<div class="cover"><div class="brand">{text}</div>'); first=False; in_cover=True; continue
        if st=='Title': parts.append(f'<div class="cover"><h1>{text}</h1>'); in_cover=True
        elif st=='Heading 1':
            close_list()
            if in_cover: parts.append('</div>'); in_cover=False
            parts.append(f'<h1>{text}</h1>')
        elif st=='Heading 2':
            close_list()
            parts.append(f'<h2>{text}</h2>')
        elif st=='Heading 3':
            close_list()
            parts.append(f'<h3>{text}</h3>')
        elif st.startswith('List Bullet'):
            if list_tag!='ul':
                if list_tag: parts.append(f'</{list_tag}>')
                parts.append('<ul>'); list_tag='ul'
            parts.append(f'<li>{text}</li>')
        elif st.startswith('List Number'):
            if list_tag!='ol':
                if list_tag: parts.append(f'</{list_tag}>')
                parts.append('<ol>'); list_tag='ol'
            parts.append(f'<li>{text}</li>')
        else:
            close_list()
            parts.append(f'<p>{text}</p>')
    elif child in tables:
        close_list()
        table=tables[child]; parts.append('<table>')
        for ri,row in enumerate(table.rows):
            parts.append('<tr>')
            for cell in row.cells:
                tag='th' if ri==0 else 'td'; parts.append(f'<{tag}>{escape(cell.text).replace(chr(10),"<br>")}</{tag}>')
            parts.append('</tr>')
        parts.append('</table>')
close_list()
if in_cover: parts.append('</div>')
parts.append('<footer>TrustForge Hermes｜HurricaneSoft｜完整商業化提案報告</footer></body></html>')
dst.write_text(''.join(parts),encoding='utf-8')
print(dst)

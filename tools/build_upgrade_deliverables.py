from pathlib import Path
from shutil import copy2
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
SLIDE=ROOT/'docs/competition/slide-deck'
OUT=ROOT/'outputs'

TERMS={
 'Runtime':'執行時環境：程式實際運作的地方。在 TrustForge 中，它接收分析請求、呼叫資料來源與 API、執行 Trust Kernel、保存執行紀錄並把結果回傳前端。它是 Agent 的執行引擎，不是模型本身。',
 'Bedrock':'Amazon Bedrock 是 AWS 的生成式 AI 服務平台。在 TrustForge 中只負責 Claims 抽取、語意立場分類與報告行文，不負責決定市場方向、計算 Trust Score 或批准升級。',
 'Trust Kernel':'信任計算核心：用可測試、可重現的規則處理來源可信度、交叉佐證、時效衰減、重複資訊、異常訊號、反方證據與 confidence。Bedrock 負責把資料說成人話，Trust Kernel 負責計算資料值得信多少。',
 'Evidence-native':'以證據為原生基礎：系統從設計開始就要求每個結論帶來源、時間、Claim 與溯源鏈，而不是先生成答案後才補引用。',
 'Decision Intelligence':'決策智能：不只整理資料，還協助使用者理解風險、信心、反方證據與什麼情況會改變結論。',
 'Agent':'人工智慧代理人：能依照受控流程自主蒐集資料、呼叫工具、分析、產出報告並留下紀錄的軟體。',
 'deterministic trust reasoning':'確定性信任推理：在相同輸入、規則與版本下，產生相同的信任評分與判斷，避免每次隨機變動。',
 'audit trail':'完整稽核軌跡：記錄資料、步驟、模型版本、分數、批准與錯誤，讓結論能被事後重播、查證與追責。',
 'Claim':'主張：從新聞、社群或公告中抽出的可驗證具體說法，例如某公司將推出新的加密服務。',
 'Claim Extraction':'主張抽取：從非結構化文字中找出可驗證的主張，讓後續規則能逐條評估，而不是只對整篇文章打分。',
 'Bullish':'看多、偏樂觀：語意上預期資產或市場可能上升。它是立場分類，不是價格預測。',
 'Bearish':'看空、偏悲觀：語意上預期資產或市場可能下跌。它是立場分類，不是交易建議。',
 'Neutral':'中立：資料沒有明確支持上漲或下跌方向。',
 'Stance':'立場分類：判斷一則資訊偏向 bullish、bearish 或 neutral。',
 'Deduplication':'去重：將同一事件的轉載或重複內容合併，避免把回音室誤算成多個獨立來源。',
 'Corroboration':'交叉佐證：確認同一 Claim 是否被不同且獨立的來源支持。',
 'Uncertainty':'不確定性：揭露資料不足、來源矛盾、模型限制與可能推翻結論的因素。',
 'Training backend':'訓練後端：真正讀取訓練資料、送出訓練任務、產生模型產物與回傳指標的執行模組。',
 'ModelHub':'模型治理中心：管理模型版本、訓練請求、指標、artifact、candidate 狀態與人工批准生命週期。',
 'SageMaker':'AWS 機器學習平台：從 S3 讀取訓練資料、建立 Training Job、執行校準訓練並產生模型 artifact；它不是 Bedrock，也不是 ModelHub。',
 'S3':'AWS 物件儲存服務：保存訓練資料、模型檔案、manifest 與其他 artifact。',
 'Isotonic Regression':'單調迴歸校準模型：把原始 confidence 調整成較可靠的機率，不是大型語言模型，也不是價格預測器。',
 'Calibration':'信心校準：用歷史結果檢查模型是否過度自信，並把 raw confidence 調整成較可靠的 calibrated confidence。',
 'artifact':'訓練或執行產物，例如模型檔、校準參數、報告、設定檔、版本資訊與評估指標。',
 'learning event':'學習事件：每次分析留下的可追蹤紀錄，包含來源、Claims、Trust Score、模型版本、run_id、Execution Log 與日後 outcome。它先沉澱資料，不代表立即重訓。',
 'T+1':'分析時間點後第 1 天的實際市場結果，用來事後檢查信心與方向是否合理。',
 'T+7':'分析時間點後第 7 天的實際市場結果，用來觀察中短期 outcome。',
 'T+14':'分析時間點後第 14 天的實際市場結果，用來觀察較長窗口 outcome。',
 'outcome':'事後結果標籤：等時間經過後取得的真實市場結果，用來評估與校準，而不是分析當下已知的答案。',
 'candidate':'候選版本：新訓練模型、校準器、規則或升級方案，已產生但尚未通過評估與人工批准。',
 'EC2':'AWS 虛擬伺服器：目前 TrustForge 部署後端與 Web 服務的運算主機。',
 'nginx':'Web Server 與 Reverse Proxy：接收 HTTPS、提供前端頁面，並把 API 請求轉給 Python backend。',
 'App Runner':'AWS 託管部署服務；它是舊的部署構想，不是目前 TrustForge 公開部署入口。',
 'run_id':'每次分析的唯一識別碼，用來把 Final Report、Evidence List、Execution Log、Snapshot 與模型版本串在一起。',
 'SHA-256 manifest':'檔案完整性清單：為每個檔案計算 SHA-256 雜湊值，之後可確認檔案是否被修改。',
 'Hermes Workspace':'Hermes 的主要工作介面，集中操作分析、比較、歷史、來源狀態、巨鯨活動與證據軌跡。',
 'Analyze':'分析：執行一次新的市場或資產分析。',
 'Compare':'比較：在相同快照或時間範圍下比較不同資產或分析結果。',
 'History':'歷史紀錄：查看過去的分析、報告、執行狀態與產物。',
 'Source Status':'來源狀態：查看資料來源是否正常、最新、可用與具備足夠 coverage。',
 'Whale Activity':'巨鯨活動：追蹤大型錢包或大額資金轉移，作為風險與異常訊號。',
 'Evidence Trail':'證據軌跡：把結論連回 Claim、原始來源、網址、時間戳與引用片段。',
 'Training Status':'模型訓練狀態：顯示訓練資料、任務、模型版本、指標、candidate 與人工審核進度。',
 'Budget':'預算護欄：控制 Bedrock API、token、Agent 次數、訓練與 AWS 運算成本，超限時限制或停止呼叫。',
 'Carbon':'算力與推理用量觀測：記錄 token、LLM 呼叫與可估算消耗；若沒有排放係數與查證邊界，不等同正式碳盤查。',
 'Module Telemetry':'模組遙測：記錄模組執行次數、耗時、成功率、錯誤、資料量與版本，用來觀測健康度與效能。',
 'Upgrade Queue':'升級候選佇列：集中管理待測試、待評估、待批准或待回滾的模型、規則與模組版本。',
 'Memory':'記憶系統：保存過去的 run、artifact、evidence、learning event 與上下文，讓 Agent 不必每次從零開始。',
 'AgentCore':'Agent 執行框架：管理工具呼叫、任務流程、狀態與受控自主行為；TrustForge 的 domain logic 仍由自身 contract 與 policy 控制。',
 'Admin Audit':'管理員稽核：記錄設定變更、模型審核、升級批准、回滾與權限操作。',
 'Comparison Snapshot':'比較快照：在特定時間固定保存資產、資料來源、Trust Score、模型版本與比較結果，之後可一致重現。',
 'Analysis Journey':'分析旅程：記錄從資料蒐集、驗證、評分、報告產出到交付的完整流程。',
 'Final Report':'最終分析報告：整理結論、信心、風險、限制、反方證據與可改變條件。',
 'Evidence List':'證據清單：列出每項結論對應的來源、網址、時間戳、引用片段與 Claim。',
 'Execution Log':'執行紀錄：記錄實際步驟、工具、耗時、成本、錯誤與結果。',
 'Ingestion':'資料匯入：抓取新聞、社群、鏈上與市場 API，檢查格式與時間、標準化並送入分析管線。',
 'Reasoning Pipeline':'推理處理管線：把原始資料依序轉成 Claims、立場、交叉佐證、Trust Score、uncertainty 與可追溯報告。',
 'Snapshot':'分析快照：保存當下資料、模型、規則與結果狀態，方便比較與重播。',
 'API Response':'API 回應：以結構化格式把分析結果提供給 frontend 或其他企業系統。',
 'Tool Registry':'工具註冊表：記錄 Agent 可使用工具的名稱、輸入輸出、權限、版本與安全限制。',
 'Memory plane':'記憶層：保存 run、artifact、evidence、learning event 與 replay history。',
 'Policy plane':'政策控制層：記錄哪些來源、分析、報告、評估與改善行為可以被誰使用。',
 'Run':'一次完整的分析執行。',
 'Replay history':'重播歷史：保存過去的執行流程，讓相同輸入與版本可以重現。',
 'Evaluation':'評估：用測試、benchmark、ECE、Brier、holdout 或品質指標檢查候選版本。',
 'Improvement':'改善：對資料、規則、模型或模組提出可審核的升級候選。',
 'Timeout':'逾時：系統在預定期限內沒有完成。',
 'Insufficient data':'資料不足：目前證據量或品質不足以支持可靠結論。',
 'Permission error':'權限錯誤：請求的工具、資料或 AWS 資源沒有被授權。',
 'Model anomaly':'模型異常：模型輸出、延遲、格式或信心出現不符合預期的狀態。',
 'Degraded':'降級模式：以較少功能、較少來源或較保守方式繼續服務。',
 'Abstain':'拒絕下結論：明確表示證據不足，避免猜測填補缺口。',
 'Fail-closed':'安全失敗：遇到高風險錯誤時停止輸出，而不是產生未驗證答案。',
 'Evidence contract':'證據契約：要求每筆資料帶 source、fetched_at、content_reference、related_claim 與可回查資訊。',
 'Source':'資料來源，例如官方公告、行情 API、新聞或鏈上資料。',
 'fetched_at':'資料被系統抓取或驗證的時間。',
 'content_reference':'原文內容的位置、引用片段或可定位的內容參照。',
 'related_claim':'與這筆 Evidence 對應的主張識別。',
 'URL':'網頁網址，用來回到原始來源。',
 'Endpoint':'API 服務端點，系統實際呼叫的服務入口。',
 'Query':'查詢條件或參數，說明系統如何取得資料。',
 'PoC':'Proof of Concept，概念驗證，用小範圍資料確認產品可行性。',
 'Pilot':'試點導入，在真實團隊與流程中小規模運作並量測 KPI。',
 'Trust Layer':'信任治理層：位於外部資料與生成模型之間，負責證據、評分、限制與稽核。',
}

TOOLTIP_CSS='<style>.term{border-bottom:1px dotted #2e74b5;cursor:help;position:relative}.term:hover:after{content:attr(data-tip);position:absolute;z-index:50;left:0;top:1.6em;width:280px;padding:9px 11px;border-radius:8px;background:#10294a;color:#fff;font-size:13px;line-height:1.45;box-shadow:0 8px 24px rgba(0,0,0,.25);white-space:normal;text-align:left}.glossary-note{background:#eef8fb;border-left:4px solid #2e74b5;padding:10px 14px;margin:12px 0}</style>'

def tooltip_html(text):
    # Replace terms with placeholders first so terms inside tooltip explanations
    # are never recursively wrapped into nested HTML spans.
    placeholders={}
    for i,(term,tip) in enumerate(sorted(TERMS.items(),key=lambda x:-len(x[0]))):
        token=f'__TRUSTFORGE_TERM_{i}__'
        text=text.replace(term,token)
        placeholders[token]=f'<span class="term" data-tip="{tip}">{term}</span>'
    for token,html in placeholders.items(): text=text.replace(token,html)
    return text

def copy_html(src,dst,inject=''):
    text=src.read_text(encoding='utf-8')
    if '</head>' in text: text=text.replace('</head>',TOOLTIP_CSS+'</head>',1)
    if inject: text=text.replace('<body>', '<body>'+inject,1)
    # Keep markup intact and wrap visible text nodes only (never attributes/tags).
    def visible(match):
        chunk=match.group(1)
        if not chunk.strip() or 'data-tip=' in chunk: return match.group(0)
        return '>'+tooltip_html(chunk)+'<'
    text=re.sub(r'>([^<>]+)<',visible,text)
    dst.write_text(text,encoding='utf-8')

def upgrade_deck():
    src=SLIDE/'TrustForge_正式提案簡報_6分鐘.html'; dst=SLIDE/'TrustForge_正式提案簡報_6分鐘_升級版.html'
    text=src.read_text(encoding='utf-8')
    css='<style>.commercial-mini{display:grid;grid-template-columns:repeat(3,1fr);gap:1vw;margin-top:1.8vh}.commercial-mini div{padding:1vw;border:1px solid rgba(65,217,232,.55);border-radius:12px;background:rgba(65,217,232,.09)}.commercial-mini b{display:block;color:#41d9e8;font-size:clamp(12px,1vw,16px);margin-bottom:5px}.commercial-mini span{color:#b7c9dc;font-size:clamp(11px,.85vw,14px);line-height:1.35}</style>'
    text=text.replace('</head>',css+'</head>',1)
    marker='<p class="quote"'
    insert='<div class="commercial-mini"><div><b>首個買方</b><span>交易所研究、風控與合規團隊</span></div><div><b>4 週 PoC</b><span>5 幣＋4 類外部來源＋Evidence 抽查</span></div><div><b>成功 KPI</b><span>查核時間、可溯源率、報告完成時間</span></div></div>'
    text=text.replace(marker,insert+marker,1)
    dst.write_text(text,encoding='utf-8')

def upgrade_script():
    src=SLIDE/'TrustForge_正式提案講稿_6分鐘.html'; dst=SLIDE/'TrustForge_正式提案講稿_6分鐘_升級版.html'
    inject='<div class="glossary-note"><b>講者使用方式：</b>藍色虛線底線的專有名詞，滑鼠移上去會顯示一句解釋；正式上台時只需先講中文，評審追問再補英文。</div>'
    copy_html(src,dst,inject)

def upgrade_qa():
    src=SLIDE/'TrustForge_正式提案_4分鐘備詢.html'; dst=SLIDE/'TrustForge_正式提案_4分鐘備詢_升級版.html'
    inject='<section class="qa key"><h2>現場 Top 6 速答</h2><p><b>自行訓練：</b>我們訓練的是 task-specific confidence calibrator，不是另訓大型語言模型。</p><p><b>為何不用現成模型：</b>Bedrock 負責語言理解；信任分數與校準需要可重現、可稽核的 deterministic pipeline。</p><p><b>ModelHub／SageMaker：</b>ModelHub 管治理，SageMaker 執行 AWS 訓練。</p><p><b>AWS 限制：</b>生成式 foundation model 只走 Bedrock；校準器是任務模型，正式路徑可走 SageMaker。</p><p><b>來源驗證：</b>台灣監管 connector 已建立，但各來源 coverage 依 live、cache、fixture 狀態揭露；BlockTempo 仍是 planned。</p><p><b>15 分鐘失敗：</b>以 budget、cache、平行來源與 fail-closed 控制；資料不足就 abstain，不硬湊答案。</p></section>'
    copy_html(src,dst,inject)

def upgrade_report():
    src=OUT/'TrustForge_完整商業化提案報告.docx'; dst=OUT/'TrustForge_完整商業化提案報告_升級版.docx'
    doc=Document(src)
    target=next((p for p in doc.paragraphs if p.text.strip()=='摘要與產品定位'),doc.paragraphs[0])
    entries=[
      ('Heading 1','Executive Summary｜商業化摘要'),
      ('Normal','目標客戶：交易所、券商研究部、虛擬資產資訊平台，以及需要可稽核市場資料的風控與合規團隊。第一個產品切入點是 HOYA BIT 市場資訊頁的 Trust Layer 外掛。'),
      ('Normal','導入方案：第一階段以 4 週 PoC 接入 5 幣與 4 類外部來源，交付 Evidence Trail 與人工抽查報告；第二階段以 8–12 週 Pilot 接入研究流程與角色權限；第三階段再進入 3–6 個月 Production，導入 SLA、audit、budget、tenant 與 dashboard。'),
      ('Normal','KPI：查核時間下降、Evidence 可溯源率、報告完成時間、人工改稿率、失敗後恢復時間與每次 run 成本。收費採 API usage、seat 與 enterprise integration 的組合，實際價格於 PoC 後依資料量與部署模式報價。'),
      ('Normal','商業風險邊界：不提供投資建議、不承諾價格預測；live、cache、fixture 明確標示；第三方 API 依授權使用；Evidence、Log 與 artifact 不輸出秘密。'),
      ('Heading 2','競品與替代方案'),
      ('Normal','一般 RAG 找得到資料，但通常不先評估來源可信度；一般 Crypto AI 能給答案，但難以追溯；BI Dashboard 有數字卻沒有推理鏈；人工研究可信但慢且難重現。TrustForge 的差異是先建立 Claim → Evidence → Trust → Report 的可治理鏈。'),
    ]
    for style,text in reversed(entries):
        p=target.insert_paragraph_before(text); p.style=style
    # metadata
    props=doc.core_properties; props.title='TrustForge Hermes｜完整商業化提案報告（升級版）'; props.author='HurricaneSoft'; props.company='HurricaneSoft'; props.subject='Evidence-native Decision Intelligence Agent 商業化提案'
    try:
        doc.save(dst)
    except PermissionError:
        # A user may have the upgraded DOCX open in Word. Keep the existing
        # reviewed DOCX and still regenerate the HTML/PDF siblings.
        pass
    src_html=OUT/'TrustForge_完整商業化提案報告.html'; dst_html=OUT/'TrustForge_完整商業化提案報告_升級版.html'
    summary='<section class="commercial-summary"><h1>Executive Summary｜商業化摘要</h1><p><b>目標客戶：</b>交易所、券商研究部、虛擬資產資訊平台、風控與合規團隊。第一個切入點是 HOYA BIT 市場資訊頁的 Trust Layer 外掛。</p><p><b>導入方案：</b>4 週 PoC（5 幣＋4 類外部來源＋Evidence 抽查）→ 8–12 週 Pilot（研究流程與角色權限）→ 3–6 個月 Production（SLA、audit、budget、tenant、dashboard）。</p><p><b>KPI／收費：</b>查核時間、可溯源率、報告完成時間、人工改稿率、恢復時間與每次 run 成本；收費採 API usage、seat、enterprise integration 組合。</p><p><b>風險邊界：</b>不提供投資建議、不承諾價格預測；live/cache/fixture 明示；第三方 API 依授權使用；Evidence、Log、artifact 不輸出秘密。</p><h2>競品與替代方案</h2><p>一般 RAG 找得到資料但不先評估可信度；一般 Crypto AI 難稽核；BI Dashboard 沒有推理鏈；人工研究慢且難重現。TrustForge 建立 Claim → Evidence → Trust → Report 的可治理鏈。</p></section>'
    copy_html(src_html,dst_html,summary)

if __name__=='__main__':
    upgrade_deck(); upgrade_script(); upgrade_qa(); upgrade_report(); print('upgrade deliverables generated')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('outputs/TrustForge_完整商業化提案報告.docx')
BLUE = '2E74B5'; DARK = '0B2545'; LIGHT = 'F4F6F9'; GOLD = '7A5A00'; RED = '9B1C1C'; GRAY = '5B6573'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    ind = tblPr.find(qn('w:tblInd'))
    if ind is None: ind = OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i])); tcW.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cell)

def style_run(run, size=11, color='000000', bold=False, italic=False):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold = bold; run.italic = italic

def add_p(doc, text='', style=None, bold_prefix=None, align=None, after=8):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True, color=DARK)
        style_run(p.add_run(text[len(bold_prefix):]))
    else: style_run(p.add_run(text))
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.208
    style_run(p.add_run(text)); return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.208
    style_run(p.add_run(text)); return p

def add_callout(doc, label, text, fill='F4F6F9', color=DARK):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_geometry(t, [9360])
    c = t.cell(0,0); set_cell_shading(c, fill); p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(label + '  '), bold=True, color=color); style_run(p.add_run(text), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(t, widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_shading(c, 'E8EEF5'); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_run(p.add_run(h), bold=True, color=DARK)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            p=cells[i].paragraphs[0]; style_run(p.add_run(str(val))); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(t, widths); doc.add_paragraph().paragraph_format.space_after=Pt(2); return t

def h1(doc, text): doc.add_heading(text, level=1)
def h2(doc, text): doc.add_heading(text, level=2)
def h3(doc, text): doc.add_heading(text, level=3)

def configure(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei'); normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string('000000'); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333; normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,12,6),('Heading 3',12,'1F4D78',8,4)]:
        s=doc.styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for section in doc.sections:
        hp=section.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT; style_run(hp.add_run('TrustForge Hermes  |  Commercial Product Proposal'),9,GRAY)
        fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; style_run(fp.add_run('HurricaneSoft  |  Confidential proposal'),9,GRAY)

def build():
    doc=Document(); configure(doc)
    # proposal_centerpiece header pattern
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); style_run(p.add_run('HURRICANESOFT 颶風軟體'),12,GRAY,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4); style_run(p.add_run('TrustForge Hermes'),28,DARK,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); style_run(p.add_run('Evidence-native Decision Intelligence Agent'),15,GRAY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(22); style_run(p.add_run('商業化產品提案報告｜加密市場資訊的可治理信任層'),11,GRAY,True)
    add_table(doc,['文件資訊','內容','文件資訊','內容'],[
        ('產品團隊','中再參與｜HurricaneSoft','提案角色','商業化產品與技術方案'),
        ('命題場域','智慧交易／金融資訊','版本','2026-08-01'),
        ('文件用途','完整書面提案；非 6 分鐘上台講稿','讀者','金融企業、產品、工程與治理團隊')], [1800,2880,1800,2880])
    add_callout(doc,'核心主張','TrustForge 不把 AI 當成會回答問題的黑箱，而是把多源資料、確定性信任推理、模型協作、治理與稽核串成可部署、可觀測、可回復的 AI Agent 產品。','EEF8FB','0B5261')
    doc.add_page_break()

    h1(doc,'摘要與產品定位')
    add_p(doc,'TrustForge Hermes 是面向金融資訊與高風險研究場景的 evidence-native decision intelligence agent。產品將行情、新聞、社群、監管、鏈上與企業資料導入同一個 evidence contract，先將內容拆解為可驗證 Claims，再透過 deterministic trust reasoning 計算來源信譽、獨立佐證、時效衰減、重複與異常訊號，最後才讓 Amazon Bedrock 將受控的 TrustedBrief 組織為可讀報告。')
    add_p(doc,'產品交付的不是單一分數或一句預測，而是可回溯的 decision record：結論、信心、限制、反方證據、來源、時間戳、模型版本、執行紀錄與產物雜湊都能被重播與抽查。這使 TrustForge 能從競賽原型升級為可嵌入研究、風控、客服、合規與營運流程的 Trust Layer。')
    add_callout(doc,'產品邊界','TrustForge 衡量資訊的可信度、完整度與可稽核性，不把輸出包裝成投資建議，也不把 Trust Score 宣稱為漲跌勝率。','FFF8E8',GOLD)
    h2(doc,'商業化價值')
    for x in ['降低研究人員逐條核對來源與重複轉載的時間成本。','讓風控與合規團隊能抽查 claim-to-source、模型版本與批准紀錄。','以 API、artifact 與 telemetry 接入既有交易、客服、研究與稽核系統。','以可替換 provider、memory、policy、calibrator 與 backend 降低單一模型或單一雲服務的鎖定風險。']: add_bullet(doc,x)

    h1(doc,'一、問題定義與企業場景')
    add_p(doc,'加密市場的資料量不是唯一問題，真正的風險是資料的權重、時效與關聯不透明。鏈上事實、價格資料、監管公告、新聞敘事與匿名社群不應被視為同等證據；同一篇內容被多站轉載也不能被誤算成多個獨立來源。傳統 RAG 或摘要器可以快速生成流暢文字，卻不一定能回答「這個主張來自哪裡、何時取得、是否被獨立來源支持、什麼情況會推翻它」。')
    add_table(doc,['企業問題','產品回應','可驗證產物'],[
        ('研究效率','以 ingestion、claim extraction、deduplication 與 evidence grouping 整理多源訊號','Analysis Journey、Final Report'),
        ('風險溝通','並列支持與反方證據，揭露 uncertainty、limits 與 could-flip 條件','Evidence List、Trust breakdown'),
        ('治理與稽核','以 run_id、Execution Log、model version 與 manifest 串接決策','Execution Log、SHA-256 manifest'),
        ('持續改善','以 learning event 與延遲 outcome 建立校準候選，通過 gate 後才進入審核','Training contract、candidate、approval receipt')], [2300,4200,2860])

    h1(doc,'二、產品總體架構')
    add_p(doc,'TrustForge 由資料平面、推理平面、控制平面與產品介面組成。每個平面都有明確責任與可替換邊界，避免將所有決策責任集中在單一 LLM prompt。')
    h2(doc,'2.1 三層執行管線')
    add_table(doc,['層級','責任','核心模組'],[
        ('Ingestion Layer','接收、清理、標準化與標記外部資料','prices、HOYA BIT、news、social、regulatory、on-chain、whale trades、CoinGecko、CMC、DefiLlama、Etherscan'),
        ('Reasoning Layer','主張抽取、立場分類、去重、交叉佐證、信任分數與 uncertainty','pipeline、Trust Kernel、corroboration、evidence grouper、freshness、direction resolution'),
        ('Agent / Delivery Layer','受界限的工具調度、報告組裝、產物交付與 API 回應','Hermes orchestrator、Bedrock、report delivery、Final Report、Evidence List、Execution Log、Snapshot')], [1900,3500,3960])
    h2(doc,'2.2 前後端產品面')
    add_p(doc,'Frontend 以 React/Vite 與 Hermes Workspace 呈現產品能力；Backend 以 Python API、analysis flow 與 formal run runtime 執行任務。兩者透過結構化 API 與同一個 run_id 連接。')
    add_table(doc,['產品面','主要能力'],[
        ('Research Workspace','Analyze、Compare、History、Source Status、Whale Activity、Evidence Trail、信心與證據分解、facts/inferences ladder、反方證據與 report download'),
        ('Operations Console','Training Status、Upgrade Queue、Module Telemetry、Budget、Carbon、Health、成本與執行狀態'),
        ('Governance Console','Admin、AGOS、policy、memory、prompt version、approval、rollback、audit record'),
        ('Integration Surface','API response、snapshot、artifact、manifest、Final Report、Evidence List、Execution Log、Source/Config')], [2300,7060])
    h2(doc,'2.3 AWS 目前部署拓樸')
    add_p(doc,'現行公開拓樸以 EC2 t3.micro + nginx 為基礎：nginx 對外提供 HTTPS 與 frontend，Python API 收斂在 127.0.0.1:8080；Amazon Bedrock 是唯一 LLM 入口，S3 保存資料與產物，IAM Instance Role、SSM Session Manager 與 CloudWatch 支援最小權限與運維觀測。App Runner 是未採用的舊構想，不列為目前已部署能力。')
    add_table(doc,['元件','目前責任','治理重點'],[
        ('EC2 + nginx','承載 React/Vite frontend 與 Python API reverse proxy','網路邊界、TLS、127.0.0.1 API binding'),
        ('Amazon Bedrock','claim extraction、stance classification、report assembly 等語言任務','只做語言工作，不直接決定 Trust Score'),
        ('S3','訓練資料、模型 artifact、部署與 manifest 保存','IAM、雜湊、保留與版本化'),
        ('IAM / SSM / CloudWatch','無長期 SSH key 的維運、權限與指標','最小權限、審計、健康檢查')], [2100,4300,2960])

    h1(doc,'三、核心功能與獨立系統')
    h2(doc,'3.1 Data Ingestion 與資料契約')
    add_p(doc,'Ingestion 將外部資料轉成統一的 source、fetched_at、content_reference、asset、data_mode 與 freshness 欄位，並保留失敗原因與覆蓋缺口。連接器各自獨立，單一來源故障不應讓整體分析假裝完整。')
    for x in ['行情與企業資料：HOYA BIT 及價格 OHLCV，作為方向基準與 delayed outcome ground truth。','新聞與社群：RSS／API、X／Reddit 等，用於事件與情緒訊號，但不與官方資料等權。','鏈上與巨鯨：Etherscan、Whale Alert、Arkham、CoinMarketCap 等，辨識大額轉帳、錢包歸因與流動性背景。','監管與官方公告：保留正式發布時間與原文範圍，支援高風險事件追溯。','資料品質：cache、safe fetch、connector reliability、freshness、data integrity 與 SHA-256 manifest。']: add_bullet(doc,x)
    h2(doc,'3.2 Trust Kernel 與信任推理')
    add_p(doc,'Trust Kernel 是產品的 deterministic trust reasoning 核心。它在 LLM 行文之前處理可重現的規則與計算：來源先驗、交叉佐證、時效衰減、重複計數、矛盾、操縱風險、confidence 與 uncertainty。相同輸入、規則與版本應得到相同結果，便於測試、回放與稽核。')
    add_callout(doc,'重要邊界','Trust Kernel 的分數是資訊信任與完整度訊號，不是交易勝率；Bedrock 不被授權單獨產生市場方向。','EEF8FB','0B5261')
    h2(doc,'3.3 Formal Run、計分與交付')
    for x in ['formal_run_coordinator、budget reservation、idempotency 與 lease 防止重複執行。','analysis flow 將 claim extraction、stance、deduplication、corroboration、Trust Kernel 與 uncertainty 串成可觀測步驟。','每次 run 以 run_id 串接 Final Report、Evidence List、Execution Log、Snapshot、Source/Config 與 API response。','遇到 timeout、資料不足、權限錯誤或模型異常，系統進入 degraded、abstain 或 fail-closed，不硬湊答案。']: add_bullet(doc,x)
    h2(doc,'3.4 Agent Runtime、Memory 與 Policy')
    add_p(doc,'Hermes orchestrator、AgentCore adapter/runtime、tool registry、memory OS 與 policy compiler/executor 共同形成受界限的 Agent Runtime。Memory plane 保存 run、artifact、evidence、learning event 與 replay history；policy plane 記錄 source、analysis、report、evaluation 與 improvement 的實際消費者與允許行為。')
    add_p(doc,'AgentCore 在產品中是可替換的執行整合層，而不是把所有 TrustForge domain logic 外包給第三方。工具呼叫、預算、權限、錯誤與狀態仍由 TrustForge 的 contract、guards 與 audit 控制。')
    h2(doc,'3.5 Upgrade Control 與自動升級治理')
    add_p(doc,'Improvement、upgrade control、canary、activation lock、deployment evidence 與 rollback record 形成升級控制平面。自主循環可以 refresh_sources、measure_quality、replay_history、diagnose_improvement 並提出 candidate，但 candidate 必須通過 benchmark、安全、成本、相容性與人工批准，才可進入 activation。')
    add_table(doc,['狀態','系統行為','不可省略的證據'],[
        ('Candidate','建立新模型、規則或模組候選，不影響 production','metrics、artifact hash、training contract'),
        ('Review','進行 quality、security、cost、compatibility 與 adversarial review','reviewer attestation、decision record'),
        ('Activate','經人工批准後才切換，保留 canary 與 activation lock','approval、deployment evidence、receipt'),
        ('Rollback','發現異常時回到已知穩定版本','rollback decision record、health check、audit trail')], [1800,4300,3260])
    h2(doc,'3.6 可觀測性、成本與碳管理')
    for x in ['Module Telemetry：記錄模組執行次數、耗時、成功率、錯誤與版本。','Budget Guard：限制 token、Bedrock 呼叫、Agent 次數與每日 USD 使用量，超限時停止或降級。','Carbon：記錄 input/output token 與推理消耗，透過 Trust Kernel 本地計算、快取與必要時才呼叫 LLM，降低不必要算力。','Audit：以 hermes audit、execution events、artifact registry 與 authenticated ledger 支援追責與重播。']: add_bullet(doc,x)

    h1(doc,'四、模型策略：Bedrock、校準模型與訓練閉環')
    h2(doc,'4.1 為什麼不只使用現成生成模型')
    add_p(doc,'Bedrock foundation model 擅長語言理解、資訊抽取與報告組裝，但它不應直接承擔企業對信任分數、校準、資料權重與升級批准的責任。生成模型的輸出具有機率性，且會受到 prompt、上下文與資料品質影響；金融資訊產品需要的是在固定資料、規則與版本下可重現、可比較、可稽核的信任訊號。')
    for x in ['以本地 deterministic feature 與校準器保留可測試性，避免把治理責任藏在 prompt。','以自身資料與 delayed outcome 了解本產品的 confidence 是否過度自信，而不是只看通用模型 benchmark。','以可替換的 ModelHub 或 SageMaker TrainingBackend 執行訓練，兩者都產出 artifact 與指標；候選版本仍由 TrustForge upgrade control 管理批准流程。','生成式 foundation model 仍只使用 Amazon Bedrock；自行訓練的是 TrustForge 的 task-specific calibrator，不是繞過競賽規則另訓第三方 foundation model。']: add_bullet(doc,x)
    h2(doc,'4.2 Learning Event 與延遲標籤')
    add_p(doc,'每次資料抓取、分析、評分與報告完成後，系統會沉澱 learning event：包含來源、Claims、Trust Score、raw confidence、模型版本、run_id、Execution Log 與後續可對照的決策狀態。這是建立訓練資料的第一步，不代表當下立即重訓。')
    add_p(doc,'市場結果必須等待 outcome 成熟：T+1、T+7、T+14 分別代表分析時間點後第 1、7、14 天的實際結果。延遲標籤可以用來檢查信心是否合理、是否過度自信與校準是否需要修正，並避免把未來資訊洩漏回訓練資料。')
    h2(doc,'4.3 訓練流程與 backend 邊界')
    add_number(doc,'Run 完成後寫入 learning event 與 immutable evidence／execution metadata。')
    add_number(doc,'T+1/T+7/T+14 outcome label 成熟後，建立 training row，執行資料契約、版本與 label leakage 檢查。')
    add_number(doc,'以 Isotonic Regression 等校準方法產生 candidate，評估 ECE、Brier、holdout 與資料覆蓋範圍。')
    add_number(doc,'依 TRAINING_BACKEND 選擇 ModelHub 或 SageMaker 執行訓練；兩者皆產生 artifact 與指標，並將候選交回 TrustForge upgrade control 進行治理。')
    add_number(doc,'candidate 進入 review、canary、人工批准與 activation；不合格則退回或淘汰，production 版本不自動被取代。')
    add_callout(doc,'實作現況','目前 training trigger 是顯式 CLI／受控流程；程式與報告尚未宣稱存在無人值守 scheduler、每次 run 立即自動重訓或自動 activation。自動排程是後續 enhancement，人工否決權永久保留。','FFF8E8',GOLD)
    h2(doc,'4.4 ModelHub 與 SageMaker 的平行後端邊界')
    add_table(doc,['元件','角色','保存／執行內容'],[
        ('ModelHub','TrainingBackend 選項','執行 fit、產生 model artifact 與 metrics，回傳候選結果'),
        ('SageMaker AI','TrainingBackend 選項','從 S3 讀取資料、建立 Training Job、產生 model artifact 與 metrics'),
        ('TrustForge upgrade control','共同治理平面','記錄 provenance、candidate status、review、approval、activation 與 rollback'),
        ('S3','Artifact storage','training dataset、calibrator artifact、manifest、版本與完整性雜湊')], [1900,2400,5060])

    h1(doc,'五、資料與證據治理')
    h2(doc,'5.1 六類資料來源與用途')
    add_table(doc,['資料類型','範例來源','信號用途與治理'],[
        ('價格／OHLCV','HOYA BIT、prices、CoinGecko','趨勢基準、T+ outcome、時間與覆蓋範圍'),
        ('新聞','RSS／API、news connector','事件主張、發布時間、原文回溯'),
        ('社群','X／Reddit、social connector','情緒與異常訊號；權重低於官方資料'),
        ('鏈上','Etherscan、Solscan、on-chain','交易明細、合約互動、巨鯨動態'),
        ('監管','官方公告、Taiwan regulatory','政策與合規風險，保留正式來源'),
        ('市場基本面','CMC、CoinGecko、DefiLlama','市值、流動性、排名與協議背景')], [1800,2700,4860])
    h2(doc,'5.2 Evidence Contract')
    add_p(doc,'每個 Claim 都應能連回 source、fetched_at、content_reference、asset、data_mode 與 related_claim。Evidence List 不只是網址清單，而是把結論與原始證據、時間、來源家族、相互關係與限制綁定在一起。對矛盾訊號，系統保留支持與反方兩邊，不用簡單多數決掩蓋少數證據。')
    h2(doc,'5.3 四項可交付物與企業交付')
    add_table(doc,['交付物','內容','企業用途'],[
        ('Final Report','結論、Trust Score、confidence、limits、could-flip、反方證據','研究決策與管理層閱讀'),
        ('Evidence List','claim_id、來源、網址、時間、引用片段與關聯','風控、客服與合規抽查'),
        ('Execution Log','工具、步驟、時間、成本、錯誤與狀態','重播、事故調查與責任追溯'),
        ('Source / Config + Snapshot','程式、配置、資料／模型版本、manifest 與快照','可重現執行與部署驗證')], [1900,4100,3360])

    h1(doc,'六、安全、治理與營運控制')
    for x in ['秘密管理：API key、AWS credentials、資料庫密碼不進 GitHub；使用環境變數、IAM Instance Role、SSM 與 secret boundary。','權限控制：tool registry 與 policy guards 限定工具、資料來源、模型與管理操作的 scope。','Fail-closed：權限錯誤、資料不足、模型異常或預算超限時停止高風險輸出，改以 degraded 或 abstain 說明限制。','版本與稽核：每次分析與升級有 run_id、model version、artifact hash、approval receipt、deployment evidence 與 rollback record。','安全審查：升級 candidate 需通過 adversarial review、dependency／secret 檢查與 production gate；不以自動化速度取代人類批准。']: add_bullet(doc,x)
    add_callout(doc,'治理定位','TrustForge 的設計精神與 ISO/IEC 42001 的風險、可追溯性、人工監督與持續監控方向一致；本提案不宣稱已取得 ISO 42001 認證或符合性證書。','EEF8FB','0B5261')

    h1(doc,'七、商業化產品路線')
    add_table(doc,['階段','產品化目標','衡量方式'],[
        ('Phase 1：Research Trust Layer','把研究／市場資訊頁接入 Trust Layer API，提供 report、evidence、execution 與 snapshot','可溯源率、來源覆蓋、分析完成時間、人工核對節省'),
        ('Phase 2：Governed Decision Services','提供白標 API、比較快照、peer metrics、asset context 與審核工作流','API adoption、review turnaround、錯誤與回滾時間'),
        ('Phase 3：Continuous Calibration','以成熟 outcome label 驅動校準 candidate、ModelHub／SageMaker 與 canary','ECE／Brier、candidate 通過率、人工批准與 rollback 指標'),
        ('Phase 4：Enterprise Control Plane','整合 IAM、audit、budget、carbon、telemetry 與多租戶 policy','成本可預測性、合規抽查成功率、服務可用性')], [2100,4700,2560])
    h2(doc,'產品護城河')
    for x in ['Evidence contract 與 claim-to-source lineage，而不是單一 prompt。','可測試的 Trust Kernel 與可替換模型／資料／儲存 backend。','Bounded autonomous agent：能量測、診斷、提案，但不能越過人類批准。','Calibration、replay、artifact provenance 與 rollback 形成持續改善閉環。','同一組產物同時服務前端使用者、工程團隊、金融風控與稽核人員。']: add_bullet(doc,x)

    h1(doc,'八、驗收與限制')
    add_p(doc,'本產品以可驗證的工程契約作為驗收基礎，而非只以畫面是否能產生文字判斷成功。')
    add_table(doc,['驗收面','最低驗收條件'],[
        ('資料真實性','每個來源標示 data_mode、取得時間、覆蓋範圍；fixture、cache、live 不混淆'),
        ('推理可重現','固定輸入、規則與版本可得到一致 Trust／Evidence 結果'),
        ('交付完整','Final Report、Evidence List、Execution Log、Source/Config 與 manifest 可由同一 run_id 串接'),
        ('訓練治理','training row 有成熟 outcome label；candidate 有 metrics、artifact hash 與人工批准狀態'),
        ('失敗安全','逾時、資料不足、權限錯誤與模型異常會 degraded、abstain 或 fail-closed')], [2300,7060])
    h2(doc,'誠實限制')
    for x in ['TrustForge 不承諾市場方向預測或投資報酬。','部分 provider／AgentCore／DynamoDB backend 是可替換介面或支援路徑，不等於全部都是目前公開部署元件。','目前沒有宣稱每次 run 後立即自動訓練、無人值守 scheduler 或自動 activation。','操縱偵測與模型品質仍需以更多標籤、回放與實際營運資料持續驗證。']: add_bullet(doc,x)

    h1(doc,'結語')
    add_p(doc,'TrustForge Hermes 的商業價值，不是再做一個會摘要市場新聞的聊天介面，而是把企業真正需要的信任、證據、成本、權限、版本與回復能力，做成可部署的 AI Agent control plane。')
    add_callout(doc,'結論','從會回答的 AI，走向值得託付、能被抽查、能被回復的 AI Agent。','EEF8FB','0B5261')
    h1(doc,'附錄：關鍵術語')
    add_table(doc,['術語','定義'],[
        ('Evidence-native','從資料契約與證據鏈開始設計，每個結論都有可追溯依據'),
        ('Deterministic trust reasoning','固定輸入、規則與版本下可重現的信任推理'),
        ('Training backend','實際讀取訓練資料、建立訓練任務與產出 artifact 的後端'),
        ('Candidate','已產生但尚未經人工批准的模型、規則或升級候選版本'),
        ('Audit trail','可追溯誰在何時以何種資料與版本做了什麼決策'),
        ('Fail-closed','遇到不確定或錯誤時停止高風險輸出，不用猜測填補缺口')], [2700,6660])
    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)

if __name__ == '__main__': build()

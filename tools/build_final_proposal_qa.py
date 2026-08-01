from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

INK = "172033"; BLUE = "3157C8"; SKY = "12A8C7"; MINT = "21B894"
PALE = "EEF6FF"; LIGHT = "F6F8FC"; AMBER = "F59E0B"; RED = "B42318"

def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcpr.append(shd)

def cell_text(cell, value, bold=False, color=INK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(value)); r.bold = bold; r.font.name = "Microsoft JhengHei"
    r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def base(title, subtitle, label):
    d = Document(); sec = d.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(.82)
    sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.68)
    normal = d.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.16
    for name, size, color in [("Title", 29, BLUE), ("Heading 1", 18, BLUE), ("Heading 2", 13, SKY), ("Heading 3", 11, INK)]:
        s = d.styles[name]; s.font.name = "Microsoft JhengHei"; s.font.size = Pt(size)
        s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.keep_with_next = True
    h = sec.header.paragraphs[0]; h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = h.add_run("TRUSTFORGE  /  TEAM 11"); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(BLUE)
    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run("2026 雲湧智生生成式 AI 黑客松｜智慧交易・HOYA BIT"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(100,110,125)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(MINT)
    p = d.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(SKY)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("中再參與｜HurricaneSoft｜2026-07-31").font.size = Pt(9)
    return d

def callout(d, text, color=BLUE):
    t = d.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); shade(c, PALE); cell_text(c, text, True, color, 11)
    c.paragraphs[0].paragraph_format.space_before = Pt(8); c.paragraphs[0].paragraph_format.space_after = Pt(8)
    d.add_paragraph()

def bullets(d, values):
    for value in values:
        p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4); p.add_run(value)

def table(d, headers, rows, widths):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, header in enumerate(headers): shade(t.rows[0].cells[i], BLUE); cell_text(t.rows[0].cells[i], header, True, "FFFFFF", 9)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if ri % 2: shade(cells[i], LIGHT)
            cell_text(cells[i], value)
    for row in t.rows:
        for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    d.add_paragraph()

def page(d): d.add_section(WD_SECTION.NEW_PAGE)

def proposal():
    d = base("TrustForge Hermes 提案報告", "加密市場多源資訊的信任提煉 AI Agent", "解題方向｜AI 與資料應用｜AWS 架構｜Live Demo")
    callout(d, "一句話提案：我們計畫在生成式 AI 寫出市場分析之前，先對每一條主張進行來源評估、交叉佐證、時效檢查與反方保留，讓結論可回到原始證據。")
    d.add_heading("團隊基本資料", 1)
    table(d, ["欄位", "內容"], [("開發環境組別", "Team 11"), ("團隊名稱", "中再參與"), ("團隊／品牌", "HurricaneSoft"), ("隊長", "王英豪"), ("成員", "嵋婕、林子彤、王榆翔（Nicholas）"), ("命題類別", "智慧交易｜HOYA BIT"), ("作品名稱", "TrustForge Hermes（信源熔爐）")], [1.55, 5.25])
    d.add_heading("提案摘要", 1)
    d.add_paragraph("TrustForge Hermes 是我們擬議開發的加密市場多源資訊信任提煉 AI Agent。它將整合行情、新聞、鏈上與社群資料，先將內容拆成可驗證主張，再由 Trust Layer 評估來源、獨立佐證、時效與風險訊號，保留支持及反方證據，最後交由 Amazon Bedrock 組織為附引用的分析報告。提案不把價格預測或投資建議作為產品主張；目標是讓研究者知道每個結論『憑什麼相信、何時應改變判斷』。驗收目標是在 15 分鐘內，由指定題目與幣種產出 Final Report、Evidence List、Execution Log 與 Source／Config。")

    page(d); d.add_heading("一、問題與命題連結", 1)
    d.add_paragraph("加密市場資訊高速、碎片化且易受轉載、同溫層與情緒操作影響。一般生成式 AI 能整理文字，卻可能把重複來源當成多方共識，或生成無法查證的引用。HOYA BIT 命題要求多源資訊分析、證據可回溯及完整執行紀錄；TrustForge 將『信任判斷』做成獨立、可稽核的中介層。")
    table(d, ["痛點", "擬議解法", "驗收證據"], [("來源多但不等於獨立", "辨識來源與主張，計算獨立佐證", "支持／矛盾證據計數"), ("LLM 引用可能失真", "先建立 Evidence contract，再限制生成", "URL、時間、引用片段、claim_id"), ("結論缺少限制", "保留反方證據與 could_flip 條件", "信心、限制、翻轉條件"), ("Demo 時間短", "設定階段時間預算與四件式輸出", "Report／Evidence／Log／Config")], [1.55, 2.8, 2.45])
    d.add_heading("使用情境", 2)
    bullets(d, ["交易所研究與客服：快速生成可追溯的市場事件摘要。", "風險與合規：抽查引用是否存在、來源是否一致、結論是否過度延伸。", "一般使用者：在做決策前看見支持證據、反方證據與資訊缺口。"])
    d.add_heading("以官方評分標準設定建置與驗收目標", 2)
    table(d, ["評分項", "權重", "預計建置與驗收特色"], [("主題切合度", "30%", "建置多源整合、證據回溯、矛盾訊號、雙軸信心與限制說明"), ("商業應用性", "25%", "以降低研究與查核成本為目標；規劃嵌入 HOYA BIT 市場資訊頁"), ("技術可行性", "20%", "規劃 Trust Kernel、Bedrock 責任分離、獨立降級與執行紀錄"), ("創意度", "15%", "規劃來源獨立性、矛盾帳本、雙軸信心、負空間情報"), ("完成度", "10%", "驗收 15 分鐘內產出 Report、Evidence、Log 與 Source/Config"), ("AWS Kiro", "+10%", "規劃以 Spec 驅動需求、驗收、設計與實作追溯")], [1.45, .7, 4.65])

    page(d); d.add_heading("二、擬議解決方案與資料流程", 1)
    d.add_heading("六類預計資料來源", 2)
    bullets(d, ["市場行情：HOYA BIT／官方 OHLCV、成交量、波動與流動性。", "鏈上活動：Etherscan、Whale Alert、持幣與大額轉帳訊號。", "新聞媒體：公開新聞與 RSS，辨識同稿轉載與發布時間。", "官方與監管：專案公告、SEC 與臺灣監理機關等第一方資料。", "社群情緒：Reddit 等公開社群的聲量、敘事與操縱風險。", "基本面與生態：CoinGecko、DefiLlama、CMC 與鏈上生態指標。"])
    callout(d, "輸入題目＋指定幣種 → 蒐集資料 → 主張抽取 → TrustScore／交叉佐證 → Bedrock 敘事化 → 四份可稽核交付物")
    table(d, ["層級", "預計工作", "預定輸出"], [("1. Ingestion", "取得 OHLCV、新聞、社群與文件；記錄來源、取得時間、查詢條件", "標準化 Document"), ("2. Trust Kernel", "拆解 Claim；計算來源聲譽、獨立佐證、時效、風險標記；保存反方", "ScoredClaim／TrustedBrief"), ("3. Agent & Delivery", "由 Amazon Bedrock 依 TrustedBrief 組織文字，強制以 claim_id 對應證據", "Final Report 等四件")], [1.25, 3.65, 1.9])
    d.add_heading("TrustScore 的定位", 2)
    d.add_paragraph("本提案將 TrustScore 定義為『資訊完整度與可溯源性』指標，而不是幣價上漲機率。建置後將以歷史資料檢查校準能力；若無法證明預測效果，就不宣稱預測能力。產品聚焦是提升判斷品質，而非製造報酬保證。")
    d.add_heading("Evidence 最小契約", 2)
    bullets(d, ["source：來源名稱或資料提供者。", "fetched_at：取得時間。", "content_reference：引用片段、檔名、查詢或資料範圍。", "related_claim：該證據支持或反駁的主張。", "網頁附 source_url；API／CSV 附 endpoint、參數、交易對、時間範圍或檔名。"])

    page(d); d.add_heading("三、擬議 AWS 與生成式 AI 技術架構", 1)
    table(d, ["服務", "規劃用途", "預定控制措施"], [("Amazon Bedrock", "作為唯一基礎模型入口；抽取、立場分類與報告敘事", "模型 ID／用量留痕、成本護欄、離線降級"), ("Amazon EC2＋nginx", "承載 React Live Demo 與 Python API", "HTTPS、服務程序管理、健康檢查"), ("Amazon DynamoDB", "保存執行狀態、冪等及協調資料", "租約與重試控制"), ("Amazon S3", "保存交付物與版本 manifest", "SHA-256、candidate／active／previous"), ("IAM＋SSM", "提供最小權限及維運存取", "不在 Repo 儲存長效憑證"), ("CloudWatch", "提供日誌、健康及錯誤觀測", "run_id 串接執行紀錄")], [1.5, 3.15, 2.15])
    d.add_heading("生成式 AI 的預定責任邊界", 2)
    bullets(d, ["Bedrock 將負責語意抽取、有限立場分類與可讀敘事，但不得捏造證據。", "信任分數、證據關聯、限制與輸出契約將由 TrustForge pipeline 管理。", "任一模型呼叫失敗時，系統將以可說明的離線結果降級，不以假資料冒充成功。", "競賽執行環境將只使用主辦允許的 AWS 服務與基礎模型。"])
    d.add_heading("建置邊界與後續擴充", 2)
    d.add_paragraph("本提案不假設 HOYA BIT 即時 API 已可使用；只有在取得正式契約並完成連線驗證後，才會對外宣稱已整合。ModelHub 與 Amazon SageMaker 將作為可替換的受控訓練後端：兩者都能接收資料、執行訓練並交付模型產物；TrustForge 自己掌握資料 Gate、評估、版本與啟用決策。任何候選模型都必須經過評估與人工核准，不會自動升級。")
    d.add_heading("自行訓練模型的初衷與目標", 2)
    d.add_paragraph("TrustForge 預計自行訓練的是信心校準模型，而不是重新訓練大型語言模型。市場結構、消息來源與操縱手法會持續變化，因此系統將保存分析當下可見的資料，再接回 T+7 等後續真實結果，以 Isotonic Regression 將 raw confidence 校準成更接近歷史表現的 calibrated confidence。最新資料用來回答現在發生什麼；後續結果用來修正我們應該相信多少。")
    bullets(d, ["更即時：讓信心基準跟上市場、來源與風險型態的變化。", "更適合任務：針對 Evidence、來源可靠度與 TrustScore 定義最佳化，而非只依賴通用模型。", "更可解釋：保留資料版本、時間範圍、評估指標、模型產物與 SHA-256。", "形成資料護城河：累積『什麼來源在什麼情境下可靠』的時間序列經驗。", "降低平台綁定：ModelHub 與 Amazon SageMaker 共用訓練後端介面，可依環境切換。", "安全改善：候選模型須通過資料 Gate、holdout 比較與人工核准，不能自行上線。"])
    d.add_paragraph("生成式 AI 推理仍由 Amazon Bedrock 提供的 AWS 基礎模型負責；自有校準模型位於信任治理層，兩者可以同時存在且責任分離。")
    d.add_paragraph("TrustForge 的設計也將納入 AI 風險管理、可追溯性、人工監督與持續監控，作為未來導入 ISO/IEC 42001 AI 管理系統的基礎；這不代表已取得認證或完成標準符合性評估。")

    page(d); d.add_heading("四、預計 Live Demo 與驗收方式", 1)
    table(d, ["預計時間", "預計操作", "驗收證據"], [("0:00–1:00", "輸入題目與指定幣種，建立 run_id", "任務與資料模式"), ("1:00–4:00", "蒐集與標準化多源資料", "來源、時間、查詢條件"), ("4:00–7:00", "主張抽取、TrustScore、支持／矛盾歸類", "Evidence Trail"), ("7:00–10:00", "Bedrock 生成附 claim_id 的報告", "結論、信心、限制、could_flip"), ("10:00–13:00", "檢查四份輸出及引用", "Report／Evidence／Log／Config"), ("13:00–15:00", "打包與上傳、保留緩衝", "檔名、hash、完成狀態")], [1.2, 2.8, 2.8])
    d.add_heading("四份預定驗收交付物", 2)
    table(d, ["交付物", "內容"], [("① Final Report", "結論／市場判斷、關鍵依據、信心、已知限制、可能推翻條件"), ("② Evidence List", "每筆證據含 source、fetched_at、content_reference、related_claim"), ("③ Execution Log", "run_id、階段時間、狀態、模型／成本資訊、錯誤與降級"), ("④ Source／Config", "可重現程式碼、設定與版本；排除所有機密憑證")], [1.6, 5.2])

    page(d); d.add_heading("五、預期價值、差異化與成功指標", 1)
    callout(d, "別人給 HOYA BIT 一個答案；TrustForge 讓 HOYA BIT 知道，這個答案值不值得相信。")
    table(d, ["一般分析 Agent", "擬議 TrustForge"], [("直接生成答案", "將先建構證據與主張關係，再生成答案"), ("重複轉載可能被視為共識", "將計算來源獨立性與交叉佐證"), ("只顯示信心分數", "將同時顯示支持、矛盾、限制與翻轉條件"), ("錯誤難追查", "每個 claim 將可回到來源、時間與引用片段"), ("成功只看文字流暢", "將以可回溯率、完整率、執行成功率與耗時驗收")], [3.05, 3.75])
    d.add_heading("預期效益", 2)
    bullets(d, ["降低研究者逐條查核與整理多源資料的時間。", "讓 HOYA BIT 的 AI 服務從『會回答』升級為『可採信、可抽查、可交代』。", "將 Trust Layer 封裝為可整合 API，延伸至研究、客服、風控與內容治理。"])
    d.add_heading("建議 KPI", 2)
    bullets(d, ["Evidence 必填欄位完整率與來源可回溯率。", "報告 claim 與 evidence 對應率。", "15 分鐘內完成四項輸出的成功率與 P95 時間。", "失敗時的降級成功率、錯誤可診斷率與單次 Bedrock 成本。"])

    page(d); d.add_heading("六、提案繳交資訊", 1)
    table(d, ["繳交項目", "內容"], [("團隊基本資料", "參見本文件「團隊基本資料」"), ("提案摘要", "參見本文件「提案摘要」"), ("完整提案簡報", "TrustForge_決賽6分鐘簡報.pptx"), ("Live Demo 部署網址", "建置與驗收完成後提供"), ("Live Demo 錄製影片", "完成正式驗收流程後提供"), ("GitHub（完整原始碼）", "專案建置後提供可審查的原始碼、設定與執行說明")], [2.0, 4.8])
    d.add_heading("資安與完整性要求", 2)
    bullets(d, ["交付內容不得包含 AWS Access Key、API Token、資料庫密碼、Workshop Access Code 或任何私密連結。", "機密將由環境變數、IAM Instance Role 或 SSM 管理；交付流程將執行 secret scan。", "專案根目錄將保留 .kiro 與其 specs、hooks、steering，不會整包加入 .gitignore。", "README 將包含環境設定、執行範例、benchmark 重現及資料／模型／索引版本。", "依賴鎖定檔、架構與資料流程文件、Demo 網址、影片網址及 Repo 權限將納入交付驗收。"])
    d.add_heading("風險與應變", 2)
    table(d, ["風險", "應變"], [("Bedrock 延遲或權限問題", "先做連線煙霧測試；提供離線降級並在 Log 明示"), ("外部資料源失效", "使用有時間戳的快取／fixture，標示 data mode"), ("15 分鐘逾時", "設定分階段時間預算，13 分鐘開始打包"), ("引用不可回溯", "驗收時抽查 URL、引用片段與 claim_id")], [2.15, 4.65])
    path = OUT / "TrustForge_賽前提案報告.docx"; d.save(path); return path

def qa():
    d = base("TrustForge 決賽 4 分鐘備詢手冊", "評審問答｜直接回答、實證、界線與追問", "完整備詢版｜主答 30 秒，追問可延伸")
    callout(d, "答題公式：先用 10 秒講結論，再用 15 秒指出畫面／檔案證據，最後用 5 秒說限制。若評審追問，再展開本手冊的技術、商業與風險細節。")
    d.add_heading("四分鐘節奏", 1)
    table(d, ["時間", "任務", "原則"], [("0:00–0:20", "主答者重述問題並直接回答", "先結論，不重講整份簡報"), ("0:20–3:30", "依序回答所有問題", "每題 25–35 秒；指向實證"), ("3:30–4:00", "補充關鍵限制與價值", "收斂到可回溯、可抽查")], [1.3, 2.35, 3.15])
    d.add_heading("開場備用句", 2)
    d.add_paragraph("謝謝評審。TrustForge 的核心不是預測漲跌，而是在生成內容前建立可稽核的信任層；以下我們會直接回答，並以 Live Demo、Evidence List 或 Execution Log 作證。")
    d.add_heading("評分項目速查", 1)
    table(d, ["評分項", "評審可能追問", "回答錨點"], [("主題 30%", "為什麼不是普通 RAG？", "獨立性、矛盾、雙軸信心、可回溯"), ("商業 25%", "HOYA BIT 為什麼需要？", "降低查核時間、內容責任邊界、白標 API"), ("技術 20%", "真的能跑、能降級嗎？", "Trust Kernel、Bedrock、訓練後端、Log"), ("創意 15%", "哪個功能是你們獨有？", "信任工具、自有校準模型、受控學習"), ("完成 10%", "當場失敗怎麼辦？", "四份交付、獨立降級、可診斷失敗")], [1.05, 2.75, 3.0])

    qas = [
      ("TrustScore 準嗎？", "它衡量資訊完整度與可溯源性，不是漲跌機率。內部預測 AUC 約 0.49，因此我們不把它包裝成投資預測；評審可在 Evidence Trail 查看分項來源。", "Evidence Trail／限制欄"),
      ("和一般 RAG 有何不同？", "一般 RAG 找到文字就交給模型；TrustForge 先拆 Claim、辨識獨立來源、保留支持與矛盾，再限制 Bedrock 只能依 claim_id 敘事。", "claim_id → Evidence"),
      ("如何避免 AI 幻覺？", "我們不只靠提示詞。報告中的重要主張必須對應 Evidence contract；沒有來源就降級或標示資料不足，不以假引用補洞。", "source_url／引用片段"),
      ("如果來源互相矛盾？", "不強行平均或刪除反方。系統保留 contrarian evidence，降低信心，並列出 could_flip，讓使用者知道哪個新證據會改變判斷。", "支持／矛盾計數"),
      ("為什麼一定要 AWS？", "競賽合規且架構可治理：Bedrock 是唯一模型入口；EC2、DynamoDB、S3、IAM／SSM、CloudWatch 分別負責服務、狀態、交付、權限與觀測。", "AWS 架構頁"),
      ("Bedrock 失敗怎麼辦？", "執行會留下錯誤與模型使用狀態，並以明確標示的離線結果降級。系統不會把 fallback 冒充真實 Bedrock 成功。", "Execution Log／ledger"),
      ("15 分鐘怎麼保證完成？", "流程有階段預算：前 10 分鐘完成分析，10–13 分鐘驗證四份輸出，13 分鐘開始打包；外部來源逾時可切換有時間戳的快取。", "階段時間戳"),
      ("資料真的來自 HOYA BIT 嗎？", "只有在取得正式 API 契約並成功連線後才會這樣宣稱；否則畫面明確顯示資料模式與來源，不以 stub 冒充企業資料。", "data mode／source"),
      ("你們的創新是什麼？", "不是多一個聊天框，而是把『信任』做成可計算、可追溯、可反駁的中介層，讓每一段生成內容都有證據路徑。", "Trust Layer 流程"),
      ("商業價值在哪裡？", "HOYA BIT 可把它整合到研究、客服與風控：縮短人工查核時間，也降低不可解釋內容帶來的品牌與合規風險。", "KPI／整合路線"),
      ("資料來源有哪些？", "我們整合六類資料：HOYA BIT／OHLCV 市場行情、鏈上活動、公開新聞、官方與監管公告、公開社群情緒，以及 CoinGecko、DefiLlama、CMC 等基本面與生態指標。每筆都保存來源、取得時間、查詢條件與引用片段。", "Evidence List／source family"),
      ("你們有自行訓練模型嗎？", "有。我們自行訓練的是 TrustForge 信心校準模型，不是大型語言模型。它把分析當下的資料接回 T+7 等後續真實結果，再用 Isotonic Regression 校正 raw confidence；Bedrock 仍負責生成式 AI 的抽取與敘事。", "Training data／calibration artifact"),
      ("為什麼不只使用 AWS 現成模型？", "AWS 基礎模型擅長語意與生成，但不知道 TrustForge 的 Evidence、來源可靠度與信心定義。自有校準模型能針對任務修正過度自信、跟上市場變化，並累積來源可靠度資料護城河。", "calibrated confidence／來源可靠度"),
      ("ModelHub 和 SageMaker 是什麼關係？", "兩者是可替換的訓練後端，不是前後串接的固定流程。ModelHub 以 req_no 執行並回傳訓練結果；Amazon SageMaker 以 Training Job 與 S3 交付產物。TrustForge 自己掌握資料 Gate、評估、版本與啟用決策。", "TrainingBackend／job／artifact SHA-256"),
      ("模型會自動更新嗎？", "不會。新模型只會成為 candidate，automatic_apply 固定為 false；必須通過資料 Gate、holdout 比較、產物驗證與人工核准，才有資格啟用。", "candidate proposal／approval record"),
      ("目前最大的限制？", "預測能力不是產品主張、即時外部 API 受現場契約與網路影響。優先確保四份交付可回溯、Demo 可降級、未驗證能力不誇大。", "limitations／could_flip"),
      ("來源獨立性怎麼判斷？", "系統先正規化來源、比對近似文本與引用關係，再將 N 篇文章收旂為 K 個獨立集群。獨立佐證先去重再計票，避免把新聞聯播當成多方共識。資料稀疏時會顯示覆蓋不足，不強行補滿。", "Evidence 來源分組／獨立佐證數"),
      ("矛盾帳本和正反方摘要有何不同？", "矛盾帳本不只保留正反文字，還保留每條訊號的來源、原文、claim_id 與信任分項。使用者可看見鏈上、價格、新聞或社群為何互相衝突，以及系統為何比較相信某一邊。", "Cross-source Signal／Trust Breakdown"),
      ("雙軸信心為什麼比單一分數好？", "方向判斷很明確，不代表資料就完整。TrustForge 把『對目前方向的判斷』與『證據覆蓋程度』分開，所以可以誠實表達『現有訊號偏多，但鏈上覆蓋不足』，不用一個黑箱數字掉包。", "Confidence Gauge／information completeness"),
      ("什麼是負空間情報？", "我們不只報告找到的事件，也報告在已揭露的資料範圍與時間窗內沒有觀察到什麼。但措辭必須是『未觀察到』，不是『不存在』；同時列出資料源覆蓋限制，避免把沉默過度解讀成利多或利空。", "Negative-space insight／coverage note"),
      ("反事實 A/B 會不會是為了 Demo 設計的假案例？", "A/B 用同一批輸入，只改變信任層是否啟用，比較去重、信心與限制如何變化。我們會明說它是賽前離線方法對照，不把合成情境說成真實攻擊，也不佔用正式 15 分鐘執行。", "A/B 對照截圖／方法說明"),
      ("HOYA BIT 要如何整合？", "短期可先在市場資訊頁加入資訊完整度、Evidence 抽查與反轉條件；中期以 Trust Layer API 回傳結構化 Claim、Evidence 與 Log，不需替換 HOYA BIT 原有交易或內容系統。建議用四週 PoC 先驗證整合成本與使用者查核時間。", "整合路線／API 契約／PoC KPI"),
      ("如何計算成本與控制 Bedrock 預算？", "每次執行在 Execution Log 記錄模型、token、成本與耗時；執行前保留預算，接近時間或金額上限時跳過非必要複審。成本不應只用平均值承諾，會以實際題型、資料量與模型用量做 PoC 基準。", "Execution Log／budget guard"),
      ("資安與機密資料如何保護？", "長效憑證不進 Git，由 IAM Instance Role 與 SSM 管理；Python API 只聽 loopback，對外由 HTTPS 與 nginx 承接。公開的 Evidence 與 Log 應採 allowlist，只顯示評審查核需要的來源、時間、模型與成本，不透出 secret 或內部錯誤堆疊。", "IAM／SSM／HTTPS／public allowlist"),
      ("Kiro 加分不只是有安裝 IDE 嗎？", "我們保留 Kiro spec 作為需求、驗收條件、設計與實作的追溯鏈，並以 hooks 與 steering 固定開發約束。評審可直接從 .kiro 記錄對照功能與驗收，而不是只看一張 IDE 截圖。", ".kiro/specs／hooks／steering"),
    ]
    page(d); d.add_heading("高機率題庫", 1)
    for idx, (q, a, proof) in enumerate(qas, 1):
        d.add_heading(f"{idx}. {q}", 2)
        p = d.add_paragraph(); r = p.add_run("回答｜"); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE); p.add_run(a)
        p = d.add_paragraph(); r = p.add_run("指證｜"); r.bold = True; r.font.color.rgb = RGBColor.from_string(MINT); p.add_run(proof)
        if idx in (4, 8): page(d)

    page(d); d.add_heading("備詢作戰表", 1)
    table(d, ["情境", "處理方式"], [("不知道確切答案", "明說目前未觀測／未驗證，再說可如何用 Log 或測試確認"), ("評審質疑準確率", "重申 TrustScore 非預測分數；展示可回溯與限制"), ("Demo 當下失敗", "打開 Execution Log，說明失敗節點並展示降級輸出"), ("追問資料真實性", "直接開 Evidence List，抽查 URL、時間與引用片段"), ("追問資安", "說明無憑證入庫、IAM／SSM、secret scan 與 .kiro 保留政策")], [2.0, 4.8])
    d.add_heading("禁止誇大的說法", 1)
    table(d, ["不要說", "改成"], [("TrustScore 能預測漲跌", "TrustScore 衡量完整度、可溯源性與佐證品質"), ("每次都使用 Bedrock", "該次是否使用由 run ledger 證明"), ("HOYA BIT API 已串好", "僅在正式契約與連線驗證完成後宣稱"), ("系統持續自我進化", "升級候選需測試、審查、核准才可啟用"), ("所有資料都是真實即時", "畫面明示 live／cache／fixture data mode")], [2.7, 4.1])
    d.add_heading("最後 20 秒收尾", 1)
    callout(d, "TrustForge 不替使用者做投資決定；它讓使用者看得見每個結論的證據、矛盾與限制。對 HOYA BIT 而言，這是把生成式 AI 從『會回答』推進到『值得採信、可以稽核』。", MINT)
    path = OUT / "TrustForge_決賽4分鐘備詢.docx"; d.save(path); return path

if __name__ == "__main__":
    print(proposal())
    print(qa())
